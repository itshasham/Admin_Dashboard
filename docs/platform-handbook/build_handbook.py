from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "NEES-Medical-Group-Platform-Handbook.docx"
SOURCE_FILES = [
    ROOT / "01-platform-overview.md",
    ROOT / "02-public-websites.md",
    ROOT / "03-admin-workspace.md",
    ROOT / "04-people-places-assets.md",
    ROOT / "05-commerce-orders-payments.md",
    ROOT / "06-catalog-content-events.md",
    ROOT / "07-security-data-api.md",
    ROOT / "08-deployment-operations.md",
    ROOT / "09-status-and-roadmap.md",
]

# compact_reference_guide preset, with editorial_cover first-page pattern.
FONT = "Calibri"
NAVY = RGBColor(0x0B, 0x25, 0x45)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x65, 0x75)
GOLD = RGBColor(0x7A, 0x5A, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F2F4F7"
BORDER = "C8D2DE"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), BORDER)

    old_grid = table._tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        old_grid.append(col)

    for row in table.rows:
        row_properties = row._tr.get_or_add_trPr()
        cant_split = row_properties.find(qn("w:cantSplit"))
        if cant_split is None:
            row_properties.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table_widths(column_count):
    patterns = {
        2: [2700, 6660],
        3: [1800, 2700, 4860],
        4: [1500, 2100, 2400, 3360],
        5: [900, 1500, 1900, 2360, 2700],
    }
    if column_count in patterns:
        return patterns[column_count]
    base = CONTENT_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_numbering(doc, num_fmt, level_text):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), level_text)
    level.append(text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    level.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(number)


def set_style_font(style, size, color, bold=False):
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, 11, RGBColor(0x1F, 0x29, 0x37), False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    configs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in configs.items():
        style = doc.styles[name]
        set_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Code Block" not in [style.name for style in doc.styles]:
        code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = doc.styles["Code Block"]
    set_style_font(code, 8.7, NAVY, False)
    code.font.name = "Consolas"
    code._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
    code._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.05


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    set_run_font(run, size=9, color=MUTED)


def configure_header_footer(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("NEES MEDICAL GROUP  /  PLATFORM HANDBOOK")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run("INTERNAL REFERENCE  ·  ")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    add_page_field(paragraph)


def add_inline(paragraph, text, size=None, color=None):
    token_pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))")
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=(size or 11) - 0.5, color=color)
        else:
            label = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token).group(1)
            run = paragraph.add_run(label)
            set_run_font(run, size=size, color=BLUE)
            run.underline = True
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph(style="Code Block")
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_FILL)
    p_pr.append(shading)
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=8.7, color=NAVY)


def mermaid_node_label(token):
    token = token.strip()
    identifier = re.match(r"^([A-Za-z][A-Za-z0-9_]*)", token)
    node_id = identifier.group(1) if identifier else token
    quoted = re.search(r'"([^"]+)"', token)
    return node_id, quoted.group(1) if quoted else None


def add_mermaid_flow(doc, lines):
    connections = []
    labels = {}
    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith(("flowchart ", "graph ")) or "-->" not in stripped:
            continue
        source_token, target_token = stripped.split("-->", 1)
        source_id, source_label = mermaid_node_label(source_token)
        target_id, target_label = mermaid_node_label(target_token)
        if source_label:
            labels[source_id] = source_label
        if target_label:
            labels[target_id] = target_label
        connections.append((source_id, target_id))

    if not connections:
        add_code_block(doc, lines)
        return

    rows = [["Source", "Connects to"]]
    rows.extend([[labels.get(source, source), labels.get(target, target)] for source, target in connections])
    add_table(doc, rows)


def split_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(doc, rows):
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=column_count)
    set_table_geometry(table, table_widths(column_count))

    for row_index, row_data in enumerate(normalized):
        row = table.rows[row_index]
        if row_index == 0:
            row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        for column_index, value in enumerate(row_data):
            cell = row.cells[column_index]
            if row_index == 0:
                shade_cell(cell, TABLE_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            add_inline(paragraph, value, size=9.2, color=NAVY if row_index == 0 else None)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_markdown(doc, text, bullet_num_id, first_section=False):
    lines = text.splitlines()
    index = 0
    in_code = False
    code_lines = []
    code_language = ""
    current_decimal_num_id = None
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if not numbered:
            current_decimal_num_id = None

        if stripped.startswith("```"):
            if in_code:
                if code_language == "mermaid":
                    add_mermaid_flow(doc, code_lines)
                else:
                    add_code_block(doc, code_lines)
                code_lines = []
                code_language = ""
                in_code = False
            else:
                in_code = True
                code_language = stripped[3:].strip().lower()
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines):
            separator = lines[index + 1].strip()
            if separator.startswith("|") and re.fullmatch(r"[|:\- ]+", separator):
                rows = [split_table_row(line)]
                index += 2
                while index < len(lines) and lines[index].strip().startswith("|"):
                    rows.append(split_table_row(lines[index]))
                    index += 1
                add_table(doc, rows)
                continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            if level == 1:
                if not first_section:
                    paragraph.paragraph_format.page_break_before = True
                first_section = False
            add_inline(paragraph, heading_match.group(2))
            index += 1
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph()
            apply_numbering(paragraph, bullet_num_id)
            add_inline(paragraph, stripped[2:])
            index += 1
            continue

        if numbered:
            if current_decimal_num_id is None:
                current_decimal_num_id = add_numbering(doc, "decimal", "%1.")
            paragraph = doc.add_paragraph()
            apply_numbering(paragraph, current_decimal_num_id)
            add_inline(paragraph, numbered.group(1))
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        paragraph = doc.add_paragraph()
        add_inline(paragraph, stripped)
        index += 1


def add_cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(118)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("PLATFORM & OPERATIONS HANDBOOK")
    set_run_font(run, size=10.5, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("NEES Medical Group")
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("Complete Business, Feature, Security, and Technical Reference")
    set_run_font(run, size=15, color=DARK_BLUE)

    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.space_after = Pt(76)
    run = description.add_run(
        "Storefront · Clinic · Admin Workspace · Backend · People · Places · Assets · Orders · Content · Events"
    )
    set_run_font(run, size=10.5, color=MUTED, italic=True)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_after = Pt(5)
    run = date.add_run("Implementation edition · 2 August 2026")
    set_run_font(run, size=11, color=NAVY, bold=True)

    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = scope.add_run("Internal reference based on the current local source implementation")
    set_run_font(run, size=9.5, color=MUTED)

    doc.add_page_break()


def add_contents(doc, decimal_num_id):
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("Contents")
    titles = [
        "Platform Overview",
        "Public Websites and Customer Journeys",
        "Admin Workspace Feature Guide",
        "People, Places, Assets, and Guest Entry",
        "Commerce, Orders, Payments, and Customers",
        "Catalog, Content, Events, and Outreach",
        "Security, Roles, Data, and API Reference",
        "Deployment and Operating Procedures",
        "Implementation Status and Roadmap",
    ]
    for title in titles:
        paragraph = doc.add_paragraph()
        apply_numbering(paragraph, decimal_num_id)
        add_inline(paragraph, title)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(14)
    note.paragraph_format.space_after = Pt(8)
    run = note.add_run("DOCUMENT STATUS")
    set_run_font(run, size=9, color=GOLD, bold=True)
    paragraph = doc.add_paragraph()
    add_inline(
        paragraph,
        "Verified against the local working copy on 2 August 2026. Production may differ where a local change has not yet been committed and deployed.",
    )


def build():
    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        configure_header_footer(section)

    bullet_num_id = add_numbering(doc, "bullet", "•")
    decimal_num_id = add_numbering(doc, "decimal", "%1.")

    add_cover(doc)
    add_contents(doc, decimal_num_id)

    for source in SOURCE_FILES:
        add_markdown(
            doc,
            source.read_text(encoding="utf-8"),
            bullet_num_id,
        )

    properties = doc.core_properties
    properties.title = "NEES Medical Group Platform Handbook"
    properties.subject = "Business, feature, security, data, and operations documentation"
    properties.author = "NEES Medical Group"
    properties.keywords = "NEES Medical, platform, admin, storefront, clinic, people, assets, orders"
    properties.comments = "Generated from the verified platform handbook source files."

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
